from __future__ import annotations
import time
from uuid import UUID
from schemas.tool import ToolCall, ToolResult
from thrall.tools.filesystem._resolve import resolve, is_protected

_MAX_CHARS = 50_000


def execute(call: ToolCall) -> ToolResult:
    start = time.monotonic()
    path = resolve(call.args.get("path", ""))

    if is_protected(path) or not path.exists():
        return _result(call.id, error=f"path not found: {path}", start=start)
    if not path.is_file():
        return _result(call.id, error=f"not a file: {path}", start=start)

    try:
        from docx import Document
    except ImportError:
        return _result(call.id, error="python-docx is not installed — run: pip install python-docx", start=start)

    max_chars = int(call.args.get("max_chars", _MAX_CHARS))
    include_tables = call.args.get("include_tables", True)

    try:
        doc = Document(str(path))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if style.startswith("Heading"):
                    parts.append(f"\n## {text}")
                else:
                    parts.append(text)

        if include_tables:
            for i, table in enumerate(doc.tables, 1):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    parts.append(f"\n[Table {i}]\n" + "\n".join(rows))

        output = "\n".join(parts)
        if len(output) > max_chars:
            output = output[:max_chars] + f"\n\n[truncated at {max_chars} chars]"

        header = f"[DOCX: {path.name} | {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables]\n\n"
        return _result(call.id, output=header + output, start=start)

    except PermissionError:
        return _result(call.id, error=f"permission denied: {path}", start=start)
    except Exception as e:
        return _result(call.id, error=str(e), start=start)


def _result(call_id: UUID, start: float, output: str | None = None, error: str | None = None) -> ToolResult:
    return ToolResult(call_id=call_id, output=output, error=error, duration_ms=int((time.monotonic() - start) * 1000))


NAME = "documents.read_docx"
DESCRIPTION = "Extract text from a Word document (.docx). Returns paragraphs with heading structure and optional table content."
PARAMETERS = {
    "path":           {"type": "string",  "required": True},
    "max_chars":      {"type": "integer", "required": False, "default": 50000},
    "include_tables": {"type": "boolean", "required": False, "default": True},
}
